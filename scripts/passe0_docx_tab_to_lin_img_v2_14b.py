# ==============================================================================
# Script : passe0_docx_tab_to_lin_img.py
# Objectif : insérer les images directement depuis les cellules, en respectant leur ordre réel
# Date : 2025-05-02
# ==============================================================================

import os
import hashlib
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from datetime import datetime
from collections import defaultdict

#Étape 1. Initialisation des chemins et constante

SCRIPT_VERSION = "v2.14"

base_dir = Path(__file__).resolve().parent.parent
DATA_DIR = base_dir / "data"
INPUT_DIR = DATA_DIR / "00docx_tab"
OUTPUT_DIR = DATA_DIR / "01docx_lin_in"
LOG_DIR = OUTPUT_DIR / "log"
IMAGE_DIR = OUTPUT_DIR / "images"

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
LOG_DIR.mkdir(parents=True, exist_ok=True)
IMAGE_DIR.mkdir(parents=True, exist_ok=True)

files = list(INPUT_DIR.glob("*.docx"))

def generer_entete(langue, id_code, nom_fichier, version_script="unknown", version="", date="", auteur=""):
    """
    Génère l'en-tête standard d'un fichier de sortie structuré.

    - langue : "fr" ou "de"
    - id_code : ex. "00-2-04"
    - nom_fichier : nom original du .docx
    - version_script : ex. "v2.14"
    - version, date, auteur : extraits du tableau (ou vides si manquants)
    """
    now = datetime.now().strftime('%Y-%m-%d %H:%M')
    return [
        "##Identification",
        f"#Script : {version_script}.py",
        f"#Run at : {now}",
        f"#ID file : {nom_fichier}",
        f"##LANG-{langue}",
        f"#ID : {id_code}",
        f"#Version : {version}",
        f"#Date : {date}",
        f"#Author : {auteur}",
        "",
        "##Introduction",
        "",
        "##Work Start"
    ]


def cell_contient_image(cell):
    for paragraph in cell._element.xpath(".//w:drawing | .//w:pict"):
        return True
    return False

#Étape 2. Extraction des images depuis les cellules
# appelé par la fonction "transformer_tableau_et_images"

def extraire_images_des_cellules(doc, image_dir, prefix):
    images = []
    images_par_cellule = {}
    hash_dict = {}
    index = 1
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for i, cell in enumerate(row.cells[:2]):
                cle = (table_idx + 1, row_idx + 1, 'Q' if i == 0 else 'R')
                for drawing in cell._element.xpath(".//w:drawing"):
                    blips = drawing.xpath(".//*[local-name()='blip']")
                    for blip in blips:
                        rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                        if rId:
                            rel = doc.part.related_parts.get(rId)
                            if rel:
                                image_data = rel.blob
                                image_hash = hashlib.sha1(image_data).hexdigest()
                                if image_hash in hash_dict:
                                    image_name = hash_dict[image_hash]
                                    print(f"[DEBUG] Image déjà connue, réutilisée : {image_name} (cellule {cle})")
                                else:
                                    image_name = f"{prefix}_image_{index:03d}_{cle[2]}_T{cle[0]}R{cle[1]}.jpeg"
                                    image_path = image_dir / image_name
                                    with open(image_path, "wb") as f:
                                        f.write(image_data)
                                    hash_dict[image_hash] = image_name
                                    index += 1
#                                    print(f"[DEBUG] Image extraite : {image_name} (table {cle[0]}, ligne {cle[1]}, cellule {cle[2]})")
                                images.append((image_name, image_dir / image_name))
                                images_par_cellule.setdefault(cle, []).append(image_name)
                                print(f"[DEBUG] dict: mémorisé  {cle} : ['{image_name}']")

    print(f"[DEBUG] Total images enregistrées (sans doublons) : {len(hash_dict)}")
    print("[DEBUG] Dictionnaire images_par_cellule :")
    for cle, noms in images_par_cellule.items():
        print(f"  {cle} : {noms}")

#    return images
    return images_par_cellule

def extraire_et_inserer_refbak(q_texte, a_texte, prefix_branche, erreurs, avertissements):
    """
    Recherche les balises #BAK ... BAK# dans q_texte et a_texte,
    extrait les références BAK (avec préfixe branche), les supprime du texte,
    et retourne (liste_refs, texte_q_nettoyé, texte_a_nettoyé)
    """
    refs = []
    balise_ouverte = False
    pattern = r"#BAK(.*?)BAK#"
    
    print("[TRACE-BAK-IN] source :")
    print("    [Q] ", q_texte.replace("\n", "⏎"))
    print("    [A] ", a_texte.replace("\n", "⏎"))
    for source, label in [(q_texte, "Q"), (a_texte, "A")]:
        matches = re.findall(pattern, source)
        if matches:
            for match in matches:
                cleaned = match.strip()
                if cleaned:
#                    ref_bak = f"BAK {prefix_branche}.{cleaned.split()[0]}"
                    ref_bak = f"BAK {prefix_branche} [{cleaned}]"

                    refs.append(ref_bak)
            # nettoyage du texte
            source_cleaned = re.sub(pattern, "", source).strip()
        else:
            source_cleaned = source
            # vérifier les balises ouvertes non fermées
            if "#BAK" in source and "BAK#" not in source:
                erreurs.append(f"[ERROR] Balise #BAK sans fermeture dans le texte {label} : {source[:80]}...")
            elif "BAK" in source:
                avertissements.append(f"[WARN] Mention BAK sans balise dans le texte {label} : {source[:80]}...")

        if label == "Q":
            q_clean = source_cleaned
        else:
            a_clean = source_cleaned

    refs_finales = sorted(set(refs))
    print(f"[TRACE-BAK] Références extraites : {refs_finales}")
    return refs_finales, q_clean.strip(), a_clean.strip()

def verifier_et_inserer_introduction(tableau, langue, lignes_out, erreurs, avertissements):
    """
    Vérifie que la ligne 2 du tableau contient #Introduction.
    Ajoute le texte dans lignes_out si présent. Loggue erreurs/warnings sinon.
    """
    try:
        ligne_intro = tableau.rows[1]
        contenu_gauche = ligne_intro.cells[0].text.strip()
        contenu_droit = ligne_intro.cells[1].text.strip()

        if contenu_gauche != "#Introduction":
            erreurs.append(f"[ERROR] Introduction absente dans la 2e ligne pour la langue {langue}.")
        else:
            if not contenu_droit:
                avertissements.append(f"[WARN] Introduction vide pour la langue {langue}.")
            lignes_out.append("")
            lignes_out.append(contenu_droit)
            lignes_out.append("")
    except Exception as e:
        erreurs.append(f"[ERROR] Impossible de lire la ligne d’introduction pour la langue {langue} : {e}")


def transformer_tableau_et_images(doc_path):
    doc = Document(doc_path)
    nom_fichier = doc_path.name
    prefix = nom_fichier.replace('.docx', '')
    id_code = '-'.join(nom_fichier.replace('.docx', '').split('-')[:3])

#recherche des données du questionnaire
    version = ""
    date = ""
    auteur = ""
    avertissements = []
    erreurs = []

    try:
        tableau = doc.tables[0]
        lignes = tableau.rows
        if len(lignes) >= 1:
            cellule_infos = lignes[0].cells[1].text.strip()
            for line in cellule_infos.splitlines():
                if "#Version" in line:
                    version = line.split(":", 1)[-1].strip()
                elif "#Date" in line:
                    date = line.split(":", 1)[-1].strip()
                elif "#Author" in line:
                    auteur = line.split(":", 1)[-1].strip()

            if not version:
                avertissements.append("[WARN] Champ #Version manquant dans la première ligne.")
            if not date:
                avertissements.append("[WARN] Champ #Date manquant dans la première ligne.")
            if not auteur:
                avertissements.append("[WARN] Champ #Author manquant dans la première ligne.")
        else:
            avertissements.append("[WARN] Tableau trop court pour contenir les champs d'en-tête.")
    except Exception as e:
        erreurs.append(f"[ERROR] Exception lors de la lecture de l'en-tête : {e}")



    entete_inseree = False
    bloc_langue_actif = True
    langues_traitees = set()
    langue = None
    lignes_out = []
    branche_bak = prefix.split("-")[0]  # ex: '40' depuis '40-2-04'
    print(f"[DEBUG] Branche BAK détectée depuis nom de fichier : {branche_bak}")

#Extraction des images du tableau
    images_par_cellule = extraire_images_des_cellules(doc, IMAGE_DIR, prefix)

# pour chaque table
    for table_idx, table in enumerate(doc.tables):
        print(f"[DEBUG] → Tableau {table_idx+1} : {len(table.rows)} lignes, {len(table.columns)} colonnes")
       
#Étape 3 Transformation du tableau en texte linéaire

        total_rows = len(table.rows)

# ========== Détection de langue à partir de la ligne 0 ==========
        ligne_0 = " | ".join([cell.text.strip() for cell in table.rows[0].cells])

        if ligne_0.startswith("Défi") and 'fr' not in langues_traitees:
            langue = 'fr'
            lignes_out.extend(generer_entete(langue, id_code, nom_fichier, SCRIPT_VERSION, version, date, auteur))
            verifier_et_inserer_introduction(tableau, langue, lignes_out, erreurs, avertissements)
            entete_inseree = True
            bloc_langue_actif = True
            langues_traitees.add(langue)
            print("[DEBUG] Bloc français détecté par la ligne 0")

        elif ligne_0.startswith("Herausforderung") and 'de' not in langues_traitees:
            langue = 'de'
            lignes_out.extend(generer_entete(langue, id_code, nom_fichier, SCRIPT_VERSION, version, date, auteur))
            verifier_et_inserer_introduction(tableau, langue, lignes_out, erreurs, avertissements)
            entete_inseree = True
            bloc_langue_actif = True
            langues_traitees.add(langue)
            print("[DEBUG] Bloc allemand détecté par la ligne 0")

        for row_idx, row in enumerate(table.rows):

    # On ignore systématiquement la première ligne (ligne 0) et la dernière (ligne n-1)
            if row_idx == 0 or row_idx == total_rows - 1:
                continue

            cellules = [cell.text.strip() for cell in row.cells]
            ligne = f"#LG{row_idx:03d} " + " | ".join(cellules).strip()
            print(f"[DEBUG] Ligne {row_idx:02d} du tableau : {ligne}")

            if not ligne and (table_idx + 1, row_idx + 1, 'Q') not in images_par_cellule and (table_idx + 1, row_idx + 1, 'R') not in images_par_cellule:
                avertissements.append(f"Ligne vide ignorée (table {table_idx+1}, ligne {row_idx+1})")
                continue

            ligne = re.sub(r'\n{2,}', '\n', ligne)  # Réduire les sauts multiples à un seul

            if not bloc_langue_actif:
                continue

            lignes_out.append(ligne)

            for type_cell in ['Q', 'R']:
                repere_cellule = f"#RC_{table_idx+1}_{row_idx+1}_{type_cell}"
                print(f"[TRACE-RC] {repere_cellule}")
                lignes_out.append(repere_cellule)

# ========== Fin de bloc : fermeture structurée ==========
        if bloc_langue_actif:
            lignes_out.append('##Work End')
            lignes_out.append('')
            if langue == 'fr':
                lignes_out.append('[PAGEBREAK]')
            bloc_langue_actif = False
            print(f"[INFO] Fin du bloc langue '{langue}' : ##Work End et éventuel saut de page ajoutés.")


    lignes_out.append('')
    lignes_out.append('##Form End')

    lignes_reformatees = []
    for ligne in lignes_out:
        if "|" in ligne and not ligne.startswith("#Q") and "#A" not in ligne:
            q_raw, a_raw = map(str.strip, ligne.split("|", 1))

            # 🔍 Vérifier si un tag #LGnnn est présent dans la question
            tag_lg = ""
            if "#LG" in q_raw and q_raw.find("#LG") < 10:
                lg_pos = q_raw.find("#LG")
                tag_lg = q_raw[lg_pos:lg_pos+7]  # ex: #LG002
                q_raw = q_raw.replace(tag_lg, "").strip()

            # 🧱 Reconstruire proprement la ligne balisée
            if tag_lg:
                ligne_balisée = f"{tag_lg} #Q####### {q_raw} | #A------- {a_raw}"
            else:
                ligne_balisée = f"#Q####### {q_raw} | #A------- {a_raw}"

            lignes_reformatees.append(ligne_balisée)

#            print(f"[AUTO] Balises Q/A ajoutées automatiquement : {ligne_balisée[:80]}...")
        else:
            lignes_reformatees.append(ligne)

    lignes_out = lignes_reformatees

    doc_out = Document()
    previous_blank = False
# Export de l'état de lignes_out avant structuration Q/A

    with open(LOG_DIR / f"{prefix}_av_struct_{SCRIPT_VERSION}.txt", "w", encoding="utf-8") as f:
        for i, ligne in enumerate(lignes_out):
            f.write(f"[{i:03d}] {ligne}\n")
    print(f"[LOG] État brut exporté dans : {LOG_DIR / f'{prefix}_av_struct_{SCRIPT_VERSION}.txt'}")

#### Étape 4. Structuration des blocs Q/A
    lignes_reformatees = []
    ligne_idx = 0
    langue_courante = "??"
    ligne_locale = 0

    while ligne_idx < len(lignes_out):
        ligne = lignes_out[ligne_idx].strip()

#récupération no de ligne dans le tableau word
        no_ligne_tab = -1
        if "#LG" in ligne and ligne.find("#LG") < 10:
            lg_pos = ligne.find("#LG")
            if ligne[lg_pos+3:lg_pos+6].isdigit():
                try:
                    no_ligne_tab = int(ligne[lg_pos+3:lg_pos+6]) + 1
                    print(f"[DEBUG ligne #LG] ligne tableau {no_ligne_tab} : {ligne}")
                except:
                    print(f"[WARN] Impossible d’extraire le numéro après #LG dans la ligne : {ligne}")

        if ligne.startswith("##LANG-"):
            langue_courante = ligne.replace("##LANG-", "").strip()
            ligne_locale = 0
            lignes_reformatees.append(ligne)
            ligne_idx += 1
            continue

        ligne_locale += 1

        # Ligne structurée de type Q | A
        if "#Q" in ligne and "|" in ligne:
            parts = ligne.split("|", 1)
            q_part = parts[0].strip()
            q_texte_raw = q_part[q_part.find("#Q") + 2:].strip()
            q_texte = q_texte_raw.lstrip("#").strip()

            a_part = parts[1].strip()
            a_texte_raw = ""
            if "#A" in a_part and "-" in a_part:
                a_texte_raw = a_part[a_part.find("#A") + 2:].lstrip("-").strip()
            else:
                a_texte_raw = a_part.strip()

            print(f"[DEBUG] q_texte vide = {not q_texte}, a_texte_raw vide = {not a_texte_raw}, no_ligne_tab = {no_ligne_tab}")

            # ✅ Si question ET réponse sont vides, on saute ce bloc
            if not q_texte and not a_texte_raw:
                ligne_ref = f"ligne tableau {no_ligne_tab}" if no_ligne_tab >= 0 else f"ligne {ligne_locale}"
                print(f"[INFO] Bloc Q/R vide ignoré ({ligne_ref})")
                avertissements.append(f"[WARN] Bloc Q/R vide ignoré (langue {langue_courante}, {ligne_ref})")
                ligne_idx += 3
                continue

            # ❌ Si une seule des deux cellules est vide → erreur
            if (not q_texte and a_texte_raw) or (q_texte and not a_texte_raw):
                if no_ligne_tab >= 0:
                    ligne_ref = f"ligne tableau {no_ligne_tab}"
                else:
                    ligne_ref = f"ligne {ligne_locale}"
                erreurs.append(f"[ERROR] Cellule vide seule (langue {langue_courante}, {ligne_ref}) — {'Q manquante' if not q_texte else 'R manquante'}")

            # Lire les repères Q/R sur les deux lignes suivantes
            repere_q = lignes_out[ligne_idx + 1].strip() if ligne_idx + 1 < len(lignes_out) else ""
            repere_r = lignes_out[ligne_idx + 2].strip() if ligne_idx + 2 < len(lignes_out) else ""

#2.12       # Extraire les références BAK et nettoyer les textes
            refs_bak, q_texte, a_texte_raw = extraire_et_inserer_refbak(q_texte, a_texte_raw, branche_bak, erreurs, avertissements)

            # Insertion de la question
            lignes_reformatees.extend([
                "", "", "#Q########################################################"])

            if repere_q.startswith("#RC_") and repere_q.endswith("_Q"):
                lignes_reformatees.append(repere_q)
                lignes_reformatees.append("@ImgSize: height_cm = 4")
                print("[DEBUG] append @ImgSize _Q")

#2.12       #insertion du texte des références BAK
            lignes_reformatees.append("")
            if refs_bak:
                lignes_reformatees.append(f"@RefBAK: {', '.join(refs_bak) if refs_bak else ''}")

            lignes_reformatees.extend(["", "", "", q_texte])

            lignes_reformatees.extend(["", "#A-----------------------------------------------------------------------------------"])

            if repere_r.startswith("#RC_") and repere_r.endswith("_R"):
                lignes_reformatees.append(repere_r)
                lignes_reformatees.append("@ImgSize: height_cm = 4")
                print("[DEBUG] append @ImgSize _R")

            lignes_reformatees.extend(["", "", a_texte_raw, ""])
            ligne_idx += 3  # Q/A + repères

        else:
            lignes_reformatees.append(ligne)
            ligne_idx += 1


    lignes_out = lignes_reformatees
  

  #Étape 5. Génération du document linéaire final (.docx)
    # === Insertion dans le document final avec images aux bons endroits ===
    ligne_idx = 0
    langue_courante = "??"
    ligne_locale = 0

    while ligne_idx < len(lignes_out):
        ligne_locale += 1
        ligne = lignes_out[ligne_idx]
        ligne_suivante = lignes_out[ligne_idx + 1] if ligne_idx + 1 < len(lignes_out) else ""

        if ligne.strip() == "[PAGEBREAK]":
            doc_out.add_page_break()
            previous_blank = False

        elif ligne.strip() == "":
            if not previous_blank:
                p = doc_out.add_paragraph('')
                p.paragraph_format.space_after = 0
                p.paragraph_format.space_before = 0
                p.paragraph_format.line_spacing = 1
                previous_blank = True

        elif ligne.strip().startswith("#Q###"):
            p = doc_out.add_paragraph(ligne.strip())
            p.paragraph_format.space_after = 0
            p.paragraph_format.space_before = 0
            p.paragraph_format.line_spacing = 1
            previous_blank = False

        elif ligne.strip().startswith("#A---"):
            p = doc_out.add_paragraph(ligne.strip())
            p.paragraph_format.space_after = 0
            p.paragraph_format.space_before = 0
            p.paragraph_format.line_spacing = 1
            previous_blank = False

        elif ligne.strip().startswith("@ImgSize:"):
            # Vérifier si une image a été insérée juste avant cette balise
            if ligne_idx > 0 and lignes_out[ligne_idx - 1].startswith("#RC_"):
                repere = lignes_out[ligne_idx - 1].strip()
                try:
                    _, t, r, qr = repere.split("_")
                    cle_img = (int(t), int(r), qr)
                    if cle_img in images_par_cellule:
                        for nom_img in images_par_cellule[cle_img]:
                            image_path = IMAGE_DIR / nom_img
                            if image_path.exists():
                                p = doc_out.add_paragraph()
                                run = p.add_run()
                                run.add_picture(str(image_path), height=Inches(4 / 2.54))
                                print(f"[INFO] Image insérée pour {cle_img} : {nom_img}")
                except Exception as e:
                    print(f"[ERROR] Problème lors de l'insertion d'image pour repère {repere} : {e}")

            # Affichage visuel de la balise ImgSize
            p = doc_out.add_paragraph()
            run = p.add_run(ligne.strip())
            run.italic = True
            from docx.shared import RGBColor
            run.font.color.rgb = RGBColor(150, 150, 150)
            p.paragraph_format.space_after = 0
            p.paragraph_format.space_before = 0
            p.paragraph_format.line_spacing = 1
            previous_blank = False

        else:
            p = doc_out.add_paragraph(ligne.strip())
            p.paragraph_format.space_after = 0
            p.paragraph_format.space_before = 0
            p.paragraph_format.line_spacing = 1
            previous_blank = False

        ligne_idx += 1
    
    output_path = OUTPUT_DIR / nom_fichier
    doc_out.save(output_path)

    from datetime import datetime

    log_file = LOG_DIR / f"{prefix}_passe0_{SCRIPT_VERSION}.log"
    with open(log_file, "w", encoding="utf-8") as log:
        log.write(f"# Log de traitement — passe0 (version : {SCRIPT_VERSION})\n")
        log.write(f"# Fichier     : {nom_fichier}\n")
        log.write(f"# Date        : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        log.write(f"# Résultat    : OK\n\n")

        log.write("## Langues détectées : " + ", ".join(langues_traitees) + "\n\n")
        log.write("Langue   | Questions | Images\n")
        log.write("-------- | ----------|--------\n")
        nb_q_total = sum(1 for l in lignes_out if "#Q###" in l)
        nb_img_total = sum(len(v) for v in images_par_cellule.values())
        for lg in langues_traitees:
            log.write(f"{lg:<8} | {nb_q_total:<10} | {nb_img_total}\n")

        if avertissements:
            log.write("\n## Avertissements :\n")
            for warn in avertissements:
                log.write(f"{warn}\n")

        if erreurs:
            log.write("\n## Erreurs :\n")
            for err in erreurs:
                log.write(f"{err}\n")

    statut_global = "NOT OK" if erreurs else ("Warning" if avertissements else "OK")
    print(f"[INFO] Fichier de sortie généré : {output_path}")
    print(f"[INFO] Log sauvegardé dans : {log_file}")
    return nom_fichier, statut_global, erreurs, avertissements

# Exécution principale
print("=== Informations sur les chemins ===")
print(f"Répertoire du script      : {base_dir}")
print(f"Dossier d'entrée          : {INPUT_DIR}")
print(f"Dossier de sortie         : {OUTPUT_DIR}")
print(f"Fichiers .docx détectés   : {len(files)}")
print("====================================\n")

global_log = []

for file in files:
    try:
        nom_fichier, statut, erreurs, avertissements = transformer_tableau_et_images(file)
        global_log.append((file.name, statut))
        print(f"{file.name} : {statut}")
    except Exception as e:
        import traceback
        traceback.print_exc()
        global_log.append((file.name, f"ERROR: {str(e)}"))
        print(f"{file.name} : ERROR")

from datetime import datetime

global_log_path = LOG_DIR / f"log_global_passe0_{SCRIPT_VERSION}.log"
with open(global_log_path, "w", encoding="utf-8") as g:
    g.write(f"# Log global — passe0 (version : {SCRIPT_VERSION})\n")
    g.write(f"# Date : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
    g.write("Fichier traité                          | Résultat\n")
    g.write("---------------------------------------|----------\n")
    for file_name, statut in global_log:
        g.write(f"{file_name:<39} | {statut}\n")