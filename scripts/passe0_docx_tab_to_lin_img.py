# ==============================================================================
# Script : passe0_docx_tab_to_lin_img.py
# Objectif : insérer les images directement depuis les cellules, en respectant leur ordre réel
# Date : 2025-05-02
# ==============================================================================

import os
import hashlib
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from datetime import datetime

#Étape 1. Initialisation des chemins et constante

SCRIPT_VERSION = "v2.10"

base_dir = Path(__file__).resolve().parent.parent
DATA_DIR = base_dir / "data"
INPUT_DIR = DATA_DIR / "00docx_tab"
OUTPUT_DIR = DATA_DIR / "01docx_lin_in"
LOG_DIR = OUTPUT_DIR / "log"
IMAGE_DIR = OUTPUT_DIR / "images"

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
LOG_DIR.mkdir(parents=True, exist_ok=True)
IMAGE_DIR.mkdir(parents=True, exist_ok=True)

files = list(INPUT_DIR.glob("*.docx"))

def generer_entete(langue, id_code, nom_fichier):
    now = datetime.now().strftime('%Y-%m-%d %H:%M')
    return [
        '##Identification',
        f'#Script : {SCRIPT_VERSION}.py',
        f'#Run at : {now}',
        f'#ID file : {nom_fichier}',
        f'##LANG-{langue}',
        f'#ID : {id_code}',
        '#Version :',
        '#Date :',
        '#Author :',
        '',
        '##Introduction',
        '',
        '##Work Start'
    ]

def cell_contient_image(cell):
    for paragraph in cell._element.xpath(".//w:drawing | .//w:pict"):
        return True
    return False

#Étape 2. Extraction des images depuis les cellules
# appelé par la fonction "transformer_tableau_et_images"

def extraire_images_des_cellules(doc, image_dir, prefix):
    images = []
    images_par_cellule = {}
    hash_dict = {}
    index = 1
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for i, cell in enumerate(row.cells[:2]):
                cle = (table_idx + 1, row_idx + 1, 'Q' if i == 0 else 'R')
                for drawing in cell._element.xpath(".//w:drawing"):
                    blips = drawing.xpath(".//*[local-name()='blip']")
                    for blip in blips:
                        rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                        if rId:
                            rel = doc.part.related_parts.get(rId)
                            if rel:
                                image_data = rel.blob
                                image_hash = hashlib.sha1(image_data).hexdigest()
                                if image_hash in hash_dict:
                                    image_name = hash_dict[image_hash]
                                    print(f"[DEBUG] Image déjà connue, réutilisée : {image_name} (cellule {cle})")
                                else:
                                    image_name = f"{prefix}_image_{index:03d}_{cle[2]}_T{cle[0]}R{cle[1]}.jpeg"
                                    image_path = image_dir / image_name
                                    with open(image_path, "wb") as f:
                                        f.write(image_data)
                                    hash_dict[image_hash] = image_name
                                    index += 1
#                                    print(f"[DEBUG] Image extraite : {image_name} (table {cle[0]}, ligne {cle[1]}, cellule {cle[2]})")
                                images.append((image_name, image_dir / image_name))
                                images_par_cellule.setdefault(cle, []).append(image_name)
                                print(f"[DEBUG] dict: mémorisé  {cle} : ['{image_name}']")

    print(f"[DEBUG] Total images enregistrées (sans doublons) : {len(hash_dict)}")
    print("[DEBUG] Dictionnaire images_par_cellule :")
    for cle, noms in images_par_cellule.items():
        print(f"  {cle} : {noms}")

    return images

def transformer_tableau_et_images(doc_path):
    doc = Document(doc_path)
    nom_fichier = doc_path.name
    prefix = nom_fichier.replace('.docx', '')
    id_code = '-'.join(nom_fichier.replace('.docx', '').split('-')[:3])
    entete_inseree = False
    bloc_langue_actif = True
    langues_traitees = set()
    langue = None
    lignes_out = []

#Extraction des images du tableau
    images_par_cellule = extraire_images_des_cellules(doc, IMAGE_DIR, prefix)

# pour chaque table
    for table_idx, table in enumerate(doc.tables):
        print(f"[DEBUG] → Tableau {table_idx+1} : {len(table.rows)} lignes, {len(table.columns)} colonnes")
       
#Étape 3 Transformation du tableau en texte linéaire

        total_rows = len(table.rows)

# ========== Détection de langue à partir de la ligne 0 ==========
        ligne_0 = " | ".join([cell.text.strip() for cell in table.rows[0].cells])

        if ligne_0.startswith("Défi") and 'fr' not in langues_traitees:
            langue = 'fr'
            lignes_out.extend(generer_entete(langue, id_code, nom_fichier))
            entete_inseree = True
            bloc_langue_actif = True
            langues_traitees.add(langue)
            print("[DEBUG] Bloc français détecté par la ligne 0")

        elif ligne_0.startswith("Herausforderung") and 'de' not in langues_traitees:
            langue = 'de'
            lignes_out.extend(generer_entete(langue, id_code, nom_fichier))
            entete_inseree = True
            bloc_langue_actif = True
            langues_traitees.add(langue)
            print("[DEBUG] Bloc allemand détecté par la ligne 0")

        for row_idx, row in enumerate(table.rows):

    # On ignore systématiquement la première ligne (ligne 0) et la dernière (ligne n-1)
            if row_idx == 0 or row_idx == total_rows - 1:
                continue

            cellules = [cell.text.strip() for cell in row.cells]
            ligne = " | ".join(cellules).strip()
            print(f"[DEBUG] Ligne {row_idx:02d} du tableau : {ligne}")

            if not ligne and (table_idx + 1, row_idx + 1, 'Q') not in images_par_cellule and (table_idx + 1, row_idx + 1, 'R') not in images_par_cellule:
                continue

            ligne = re.sub(r'\n{2,}', '\n', ligne)  # Réduire les sauts multiples à un seul

            if not bloc_langue_actif:
                continue

            lignes_out.append(ligne)

            for type_cell in ['Q', 'R']:
                repere_cellule = f"#RC_{table_idx+1}_{row_idx+1}_{type_cell}"
                print(f"[TRACE-RC] {repere_cellule}")
                lignes_out.append(repere_cellule)

# ========== Fin de bloc : fermeture structurée ==========
        if bloc_langue_actif:
            lignes_out.append('##Work End')
            lignes_out.append('')
            if langue == 'fr':
                lignes_out.append('[PAGEBREAK]')
            bloc_langue_actif = False
            print(f"[INFO] Fin du bloc langue '{langue}' : ##Work End et éventuel saut de page ajoutés.")


    lignes_out.append('')
    lignes_out.append('##Form End')

    lignes_reformatees = []
    for ligne in lignes_out:
        if "|" in ligne and not ligne.startswith("#Q") and "#A" not in ligne:
            q_raw, a_raw = map(str.strip, ligne.split("|", 1))
            ligne_balisée = f"#Q####### {q_raw} | #A------- {a_raw}"
            lignes_reformatees.append(ligne_balisée)
            print(f"[AUTO] Balises Q/A ajoutées automatiquement : {ligne_balisée[:80]}...")
        else:
            lignes_reformatees.append(ligne)
    lignes_out = lignes_reformatees

    doc_out = Document()
    previous_blank = False
# Export de l'état de lignes_out avant structuration Q/A

    with open(LOG_DIR / f"{prefix}_av_struct_{SCRIPT_VERSION}.txt", "w", encoding="utf-8") as f:
        for i, ligne in enumerate(lignes_out):
            f.write(f"[{i:03d}] {ligne}\n")
    print(f"[LOG] État brut exporté dans : {LOG_DIR / f'{prefix}_av_struct_{SCRIPT_VERSION}.txt'}")

#### Étape 4. Structuration des blocs Q/A
    lignes_reformatees = []
    ligne_idx = 0
    while ligne_idx < len(lignes_out):
        ligne = lignes_out[ligne_idx]

        # Ligne structurée de type Q | A
        if "#Q" in ligne and "|" in ligne:
            parts = ligne.split("|", 1)
            q_part = parts[0].strip()
            q_texte_raw = q_part[q_part.find("#Q") + 2:].strip()
            q_texte = q_texte_raw.lstrip("#").strip()

            a_part = parts[1].strip()

            lignes_reformatees.extend([
                "", "", "#Q########################################################", "", "", "", q_texte
            ])

            # Lire les repères Q/R sur les deux lignes suivantes
            repere_q = lignes_out[ligne_idx + 1].strip() if ligne_idx + 1 < len(lignes_out) else ""
            repere_r = lignes_out[ligne_idx + 2].strip() if ligne_idx + 2 < len(lignes_out) else ""

            # Insérer repère Q après la question
            if repere_q.startswith("#RC_") and repere_q.endswith("_Q"):
                lignes_reformatees.append(repere_q)

            # Ligne de début de réponse
            lignes_reformatees.extend(["", "#A-----------------------------------------------------------------------------------"])

            # Insérer repère R juste après #A...
            if repere_r.startswith("#RC_") and repere_r.endswith("_R"):
                lignes_reformatees.append(repere_r)

            # Ajout du contenu de la réponse
            lignes_reformatees.extend(["", ""])
            if "#A" in a_part and "-" in a_part:
                a_texte_raw = a_part[a_part.find("#A") + 2:].lstrip("-").strip()
                lignes_reformatees.append(a_texte_raw)
            else:
                print(f"[⚠️  WARN] Ligne avec #Q mais #A malformé : {ligne}")
                lignes_reformatees.append(a_part)

            lignes_reformatees.append("")
            ligne_idx += 3  # Q/A + repères

        else:
            lignes_reformatees.append(ligne)
            ligne_idx += 1

    lignes_out = lignes_reformatees
  

  #Étape 5. Génération du document linéaire final (.docx)
    # === Insertion dans le document final avec images aux bons endroits ===
    ligne_idx = 0
    while ligne_idx < len(lignes_out):
        ligne = lignes_out[ligne_idx]
        ligne_suivante = lignes_out[ligne_idx + 1] if ligne_idx + 1 < len(lignes_out) else ""

        if ligne.strip() == "[PAGEBREAK]":
            doc_out.add_page_break()
            previous_blank = False

        elif ligne.strip() == "":
            if not previous_blank:
                p = doc_out.add_paragraph('')
                p.paragraph_format.space_after = 0
                p.paragraph_format.space_before = 0
                p.paragraph_format.line_spacing = 1
                previous_blank = True

        elif ligne.strip().startswith("#Q###"):
            p = doc_out.add_paragraph(ligne.strip())
            p.paragraph_format.space_after = 0
            p.paragraph_format.space_before = 0
            p.paragraph_format.line_spacing = 1
            previous_blank = False

        elif ligne.strip().startswith("#A---"):
            p = doc_out.add_paragraph(ligne.strip())
            p.paragraph_format.space_after = 0
            p.paragraph_format.space_before = 0
            p.paragraph_format.line_spacing = 1
            previous_blank = False


        else:
            p = doc_out.add_paragraph(ligne.strip())
            p.paragraph_format.space_after = 0
            p.paragraph_format.space_before = 0
            p.paragraph_format.line_spacing = 1
            previous_blank = False

        ligne_idx += 1
    
    output_path = OUTPUT_DIR / nom_fichier
    doc_out.save(output_path)
    return nom_fichier

# Exécution principale
print("=== Informations sur les chemins ===")
print(f"Répertoire du script      : {base_dir}")
print(f"Dossier d'entrée          : {INPUT_DIR}")
print(f"Dossier de sortie         : {OUTPUT_DIR}")
print(f"Fichiers .docx détectés   : {len(files)}")
print("====================================\n")

for file in files:
    result = transformer_tableau_et_images(file)
    print(f"{file.name} : OK")