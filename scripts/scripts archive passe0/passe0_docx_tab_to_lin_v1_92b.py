# ==============================================================================
# Script : passe0_docx_tab_to_lin_v1_92a.py
# Objectif : insérer les images directement depuis les cellules, en respectant leur ordre réel
# Date : 2025-05-02
# ==============================================================================

import os
import hashlib
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from datetime import datetime

SCRIPT_VERSION = "v1_92b"

base_dir = Path(__file__).resolve().parent.parent
DATA_DIR = base_dir / "data"
INPUT_DIR = DATA_DIR / "00docx_tab"
OUTPUT_DIR = DATA_DIR / "01docx_lin_in"
LOG_DIR = OUTPUT_DIR / "log"
IMAGE_DIR = OUTPUT_DIR / "images"

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
LOG_DIR.mkdir(parents=True, exist_ok=True)
IMAGE_DIR.mkdir(parents=True, exist_ok=True)

files = list(INPUT_DIR.glob("*.docx"))

def generer_entete(langue, id_code, nom_fichier):
    now = datetime.now().strftime('%Y-%m-%d %H:%M')
    return [
        '##Identification',
        f'#Script : {SCRIPT_VERSION}.py',
        f'#Run at : {now}',
        f'#ID file : {nom_fichier}',
        f'##LANG-{langue}',
        f'#ID : {id_code}',
        '#Version :',
        '#Date :',
        '#Author :',
        '',
        '##Introduction',
        '',
        '##Work Start'
    ]

def cell_contient_image(cell):
    for paragraph in cell._element.xpath(".//w:drawing | .//w:pict"):
        return True
    return False

def extraire_images_des_cellules(doc, image_dir, prefix):
    images = []
    images_par_cellule = {}
    hash_dict = {}
    index = 1
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for i, cell in enumerate(row.cells[:2]):
                cle = (table_idx + 1, row_idx + 1, 'Q' if i == 0 else 'R')
                for drawing in cell._element.xpath(".//w:drawing"):
                    blips = drawing.xpath(".//*[local-name()='blip']")
                    for blip in blips:
                        rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                        if rId:
                            rel = doc.part.related_parts.get(rId)
                            if rel:
                                image_data = rel.blob
                                image_hash = hashlib.sha1(image_data).hexdigest()
                                if image_hash in hash_dict:
                                    image_name = hash_dict[image_hash]
                                    print(f"[DEBUG] Image déjà connue, réutilisée : {image_name} (cellule {cle})")
                                else:
                                    image_name = f"{prefix}_image_{index:03d}_{cle[2]}_T{cle[0]}R{cle[1]}.jpeg"
                                    image_path = image_dir / image_name
                                    with open(image_path, "wb") as f:
                                        f.write(image_data)
                                    hash_dict[image_hash] = image_name
                                    index += 1
#                                    print(f"[DEBUG] Image extraite : {image_name} (table {cle[0]}, ligne {cle[1]}, cellule {cle[2]})")
                                images.append((image_name, image_dir / image_name))
                                images_par_cellule.setdefault(cle, []).append(image_name)
                                print(f"[DEBUG] dict: mémorisé  {cle} : ['{image_name}']")

    print(f"[DEBUG] Total images enregistrées (sans doublons) : {len(hash_dict)}")
    print("[DEBUG] Dictionnaire images_par_cellule :")
    for cle, noms in images_par_cellule.items():
        print(f"  {cle} : {noms}")

    return images

def transformer_tableau_et_images(doc_path):
    doc = Document(doc_path)
    nom_fichier = doc_path.name
    prefix = nom_fichier.replace('.docx', '')
    id_code = '-'.join(nom_fichier.replace('.docx', '').split('-')[:3])
    entete_inseree = False
    bloc_langue_actif = True
    langues_traitees = set()
    langue = None
    lignes_out = []
    
#    print("=== TEST POINT 1===")

    # Prétraitement Q/A : restructuration lisible
    lignes_reformatees = []
    for ligne in lignes_out:
        if "#Q" in ligne and "|" in ligne:
            parts = ligne.split("|", 1)
            q_part = parts[0].strip()
            q_texte_raw = q_part[q_part.find("#Q") + 2:].strip()
            q_texte = q_texte_raw.lstrip("#").strip()
            lignes_reformatees.append("#Q########################################################")
            lignes_reformatees.append(q_texte)
            a_part = parts[1].strip()
            if a_part.startswith("#A") and "-" in a_part:
                a_texte = a_part[a_part.find("#A") + 2:].lstrip("-").strip()
                lignes_reformatees.append("")
                lignes_reformatees.append("#A-----------------------------------------------------------------------------------")
                lignes_reformatees.append(a_texte)
            else:
                print(f"[⚠️  WARN] Ligne avec #Q mais #A malformé : {ligne}")
                lignes_reformatees.append(parts[1].strip())
        else:
            lignes_reformatees.append(ligne)
    lignes_out = lignes_reformatees

    doc_out = Document()
    previous_blank = False

    images_par_cellule = extraire_images_des_cellules(doc, IMAGE_DIR, prefix)

    cles_par_ligne = []

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            cellules = [cell.text.strip() for cell in row.cells]
            ligne = " | ".join(cellules).strip()
#            if not ligne:
#                continue
            if not ligne and (table_idx + 1, row_idx + 1, 'Q') not in images_par_cellule and (table_idx + 1, row_idx + 1, 'R') not in images_par_cellule:
                continue


            ligne = re.sub(r'\n{2,}', '\n', ligne)  # Réduire les sauts multiples à un seul

            if ligne.startswith("Défi |"):
                langue = 'fr'
                if langue in langues_traitees:
                    bloc_langue_actif = False
                    continue
                langues_traitees.add(langue)
                lignes_out.extend(generer_entete(langue, id_code, nom_fichier))
                entete_inseree = True
                bloc_langue_actif = True
                continue

            if ligne.startswith("Herausforderung |"):
                langue = 'de'
                if langue in langues_traitees:
                    bloc_langue_actif = False
                    continue
                langues_traitees.add(langue)
                lignes_out.extend(generer_entete(langue, id_code, nom_fichier))
                entete_inseree = True
                bloc_langue_actif = True
                continue

            if not bloc_langue_actif:
                continue

            if 'Temps passé à ce défi' in ligne or 'Für diese Herausforderung aufgewendete Zeit' in ligne:
                if entete_inseree:
                    lignes_out.append('##Work End')
                    lignes_out.append('')
                    if langue == 'fr':
                        lignes_out.append('[PAGEBREAK]')
                bloc_langue_actif = False
                continue

            lignes_out.append(ligne)
            cles_par_ligne.append({"cle_q": (table_idx + 1, row_idx + 1, 'Q'), "cle_r": (table_idx + 1, row_idx + 1, 'R')})
#            print(f"[TRACE] Ligne insérée (table={table_idx+1}, ligne={row_idx+1}, colonne={type_cell}): {ligne.strip()}")
#            print(f"[TRACE-L] Ligne insérée (table={table_idx+1}, ligne={row_idx+1}): {ligne.strip()}")
            for type_cell in ['Q', 'R']:
                cle = (table_idx + 1, row_idx + 1, type_cell)
                if cle in images_par_cellule:
                    lignes_out.append(f"dict:  {cle} : {images_par_cellule[cle]}")
            # Ajout du marqueur dict: à l'endroit des images (Q/R)
#            for type_cell in ['Q', 'R']:
#                cle = (table_idx + 1, row_idx + 1, type_cell)
#                print(f"[DEBUG] Test clé : {cle} → {'OK' if cle in images_par_cellule else 'absente'}")
#                if cle in images_par_cellule:
#                    lignes_out.append(f"dict:  {cle} : {images_par_cellule[cle]}")
#                    print(f"[TRACE] dict inséré dans lignes_out : {cle}")

            if table_idx == 0 and row_idx == 1:
                lignes_out.append("dict:  (1, 2, 'Q') : ['test_image_Q.jpeg']")



            for type_cell in ['Q', 'R']:
                print(f"[TRACE-LC] Ligne insérée (table={table_idx+1}, ligne={row_idx+1}, colonne={type_cell}): {ligne.strip()}")

                cle = (table_idx + 1, row_idx + 1, type_cell)
                print(f"[DEBUG] Test clé : {cle}", end=' ')
                if cle in images_par_cellule:
                    print("→ présente")
                    lignes_out.append(f"dict:  {cle} : {images_par_cellule[cle]}")
                else:
                    print("→ absente")

    lignes_out.append('')
    lignes_out.append('##Form End')

   
    
    # === Structuration des blocs Q/A ===
    lignes_reformatees = []
    ligne_idx = 0
    while ligne_idx < len(lignes_out):
        ligne = lignes_out[ligne_idx]

        # Ligne structurée de type Q | A
        if "#Q" in ligne and "|" in ligne:
            cle_q = cles_par_ligne[ligne_idx]["cle_q"]
            cle_r = cles_par_ligne[ligne_idx]["cle_r"]

            parts = ligne.split("|", 1)
            q_part = parts[0].strip()
            q_texte_raw = q_part[q_part.find("#Q") + 2:].strip()
            q_texte = q_texte_raw.lstrip("#").strip()

            a_part = parts[1].strip()

            lignes_reformatees.extend(["", "", "#Q########################################################"])
            if cle_q in images_par_cellule:
                lignes_reformatees.append(f"#DICT{cle_q} : {images_par_cellule[cle_q]}")
            lignes_reformatees.extend(["", "", q_texte, "", "#A-----------------------------------------------------------------------------------"])
            if cle_r in images_par_cellule:
                lignes_reformatees.append(f"#DICT{cle_r} : {images_par_cellule[cle_r]}")
            lignes_reformatees.append("")
            if "#A" in a_part and "-" in a_part:
                a_texte_raw = a_part[a_part.find("#A") + 2:].lstrip("-").strip()
                lignes_reformatees.append(a_texte_raw)
            else:
                print(f"[⚠️  WARN] Ligne avec #Q mais #A malformé : {ligne}")
                lignes_reformatees.append(a_part)
            lignes_reformatees.append("")
            ligne_idx += 1
        else:
            lignes_reformatees.append(ligne)
            ligne_idx += 1
# === Insertion dans le document final avec images aux bons endroits ===

    ligne_idx = 0
    while ligne_idx < len(lignes_out):
        ligne = lignes_out[ligne_idx]
        ligne_suivante = lignes_out[ligne_idx + 1] if ligne_idx + 1 < len(lignes_out) else ""

        if ligne.strip() == "[PAGEBREAK]":
            doc_out.add_page_break()
            previous_blank = False

        elif ligne.strip() == "":
            if not previous_blank:
                p = doc_out.add_paragraph('')
                p.paragraph_format.space_after = 0
                p.paragraph_format.space_before = 0
                p.paragraph_format.line_spacing = 1
                previous_blank = True

        elif ligne.strip().startswith("#Q###"):
            p = doc_out.add_paragraph(ligne.strip())
            p.paragraph_format.space_after = 0
            p.paragraph_format.space_before = 0
            p.paragraph_format.line_spacing = 1
            previous_blank = False

        elif ligne.strip().startswith("#A---"):
            p = doc_out.add_paragraph(ligne.strip())
            p.paragraph_format.space_after = 0
            p.paragraph_format.space_before = 0
            p.paragraph_format.line_spacing = 1
            previous_blank = False

        elif ligne.strip().startswith("dict:"):
            p = doc_out.add_paragraph(ligne.strip())
            p.runs[0].italic = True
            from docx.shared import RGBColor
            p.runs[0].font.color.rgb = RGBColor(150, 150, 150)
            p.paragraph_format.space_after = 0
            p.paragraph_format.space_before = 0
            p.paragraph_format.line_spacing = 1
            previous_blank = False

        else:
            p = doc_out.add_paragraph(ligne.strip())
            p.paragraph_format.space_after = 0
            p.paragraph_format.space_before = 0
            p.paragraph_format.line_spacing = 1
            previous_blank = False

        ligne_idx += 1
    
    output_path = OUTPUT_DIR / nom_fichier
    doc_out.save(output_path)
    return nom_fichier

# Exécution principale
print("=== Informations sur les chemins ===")
print(f"Répertoire du script      : {base_dir}")
print(f"Dossier d'entrée          : {INPUT_DIR}")
print(f"Dossier de sortie         : {OUTPUT_DIR}")
print(f"Fichiers .docx détectés   : {len(files)}")
print("====================================\n")

for file in files:
    result = transformer_tableau_et_images(file)
    print(f"{file.name} : OK")