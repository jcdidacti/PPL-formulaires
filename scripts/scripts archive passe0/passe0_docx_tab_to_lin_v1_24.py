
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from datetime import datetime

SCRIPT_VERSION = "v1_24"

# Répertoires dynamiques
base_dir = Path(os.path.abspath(__file__)).resolve().parent.parent
DATA_DIR = base_dir / "data"

INPUT_DIR = DATA_DIR / "00docx_tab"
OUTPUT_DIR = DATA_DIR / "01docx_lin_in"
LOG_DIR = OUTPUT_DIR / "log"

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
LOG_DIR.mkdir(parents=True, exist_ok=True)

print("=== Informations sur les chemins ===")
print(f"Répertoire du script      : {base_dir.resolve()}")
print(f"Dossier d'entrée          : {INPUT_DIR.resolve()}")
print(f"Dossier de sortie         : {OUTPUT_DIR.resolve()}")
files = list(INPUT_DIR.glob("*.docx"))
print(f"Fichiers .docx détectés   : {len(files)}")
print("====================================\n")

def decouper_balises_speciales(texte):
    texte = texte.strip()
    if '#A-' in texte:
        idx = texte.find('#A-')
        return [texte[:idx].strip(), texte[idx:].strip()]
    return [texte]

def generer_entete(langue, id_code, nom_fichier):
    now = datetime.now().strftime('%Y-%m-%d %H:%M')
    return [
        '##Identification',
        '#Script : ' + SCRIPT_VERSION + '.py',
        f'#Run at : {now}',
        f'#ID file : {nom_fichier}',
        f'##LANG-{langue}',
        f'#ID : {id_code}',
        '#Version : ',
        '#Date : ',
        '#Author : ',
        '',
        '##Introduction',
        '',
        '##Work Start'
    ]

def transformer_tableau_et_images(doc_path):
    doc = Document(doc_path)
    lignes = []
    nom_fichier = doc_path.name
    id_code = '-'.join(nom_fichier.replace('.docx', '').split('-')[:3])
    entete_inseree = False
    langue = None
    lignes_out = []

    for table in doc.tables:
        for row in table.rows:
            cellules = [cell.text.strip() for cell in row.cells]
            ligne = " | ".join(cellules).strip()
            if not ligne.strip():
                continue  # ligne vide sautée (point 3)
            if not ligne:
                continue
            if ligne.startswith('Défi |') or ligne.startswith('Herausforderung |'):
                if not entete_inseree:
                    langue = 'fr' if ligne.startswith('Défi |') else 'de'
                    lignes_out.extend(generer_entete(langue, id_code, nom_fichier))
                    entete_inseree = True
                continue
            lignes_out.extend(decouper_balises_speciales(ligne))

    if entete_inseree:
        lignes_out.append('##Work End')
        if langue == 'de':
            lignes_out.append('')
            lignes_out.append('##Form End')

    doc_out = Document()

    # Réduction des lignes vides consécutives à 1 max
    cleaned_out = []
    empty = False
    for ligne in lignes_out:
        if ligne.strip() == '':
            if not empty:
                cleaned_out.append('')
            empty = True
        else:
            cleaned_out.append(ligne)
            empty = False
    lignes_out = cleaned_out
    for l in lignes_out:
        p = doc_out.add_paragraph(l)
        p.paragraph_format.space_after = 0

    part = doc.part
    image_rels = [rel for rel in part._rels.values() if rel.reltype == RT.IMAGE]
    image_count = 0
    if image_rels:
        for rel in image_rels:
            try:
                img_data = rel.target_part.blob
                image_name = rel.target_ref.rsplit("/", 1)[-1]
                temp_path = OUTPUT_DIR / image_name
                with open(temp_path, "wb") as img_file:
                    img_file.write(img_data)
                doc_out.add_picture(str(temp_path), width=Inches(4.5))
                p = doc_out.add_paragraph(f"[image: {image_name}]")
                p.paragraph_format.space_after = 0
                image_count += 1
            except Exception as e:
                p = doc_out.add_paragraph(f"[Erreur image : {str(e)}]")
                p.paragraph_format.space_after = 0

    output_path = OUTPUT_DIR / nom_fichier
    doc_out.save(output_path)

    log_path = LOG_DIR / nom_fichier.replace(".docx", ".log")
    with log_path.open("w", encoding="utf-8") as f:
        f.write(f"# Script : {SCRIPT_VERSION}.py\n")
        f.write(f"# Version : {SCRIPT_VERSION}\n")
        f.write(f"# Date : {datetime.now().strftime('%Y-%m-%d')}\n")
        f.write("# ==========================================\n")
        f.write(f"Fichier transformé : {nom_fichier}\n")
        f.write(f"Lignes extraites : {len(lignes_out)}\n")
        f.write(f"Images extraites : {image_count}\n")

    return nom_fichier

# Exécution
log_global_path = LOG_DIR / "_global_transformation.log"
with log_global_path.open("w", encoding="utf-8") as f_global:
    f_global.write(f"# Script : {SCRIPT_VERSION}.py\n")
    f_global.write(f"# Version : {SCRIPT_VERSION}\n")
    f_global.write(f"# Date : {datetime.now().strftime('%Y-%m-%d')}\n")
    f_global.write("# ==========================================\n")
    for file in files:
        result = transformer_tableau_et_images(file)
        f_global.write(f"{result} : OK\n")
        print(f"{result} : OK")
