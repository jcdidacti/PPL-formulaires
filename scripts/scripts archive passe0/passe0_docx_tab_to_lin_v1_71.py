# ==============================================================================
# Script : passe0_docx_tab_to_lin_v1_68.py
# Objectif : remplacer la première balise #IMGQ# dans le texte par l’image extraite
# Date : 2025-05-02
# ==============================================================================

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from datetime import datetime

SCRIPT_VERSION = "v1_71"

base_dir = Path(__file__).resolve().parent.parent
DATA_DIR = base_dir / "data"
INPUT_DIR = DATA_DIR / "00docx_tab"
OUTPUT_DIR = DATA_DIR / "01docx_lin_in"
LOG_DIR = OUTPUT_DIR / "log"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
LOG_DIR.mkdir(parents=True, exist_ok=True)

files = list(INPUT_DIR.glob("*.docx"))

def generer_entete(langue, id_code, nom_fichier):
    now = datetime.now().strftime('%Y-%m-%d %H:%M')
    return [
        '##Identification',
        f'#Script : {SCRIPT_VERSION}.py',
        f'#Run at : {now}',
        f'#ID file : {nom_fichier}',
        f'##LANG-{langue}',
        f'#ID : {id_code}',
        '#Version :',
        '#Date :',
        '#Author :',
        '',
        '##Introduction',
        '',
        '##Work Start'
    ]

def decouper_balises_speciales(texte):
    texte = texte.strip()
    if '#A-' in texte:
        idx = texte.find('#A-')
        return [texte[:idx].strip(), texte[idx:].strip()]
    return [texte]

def cell_contient_image(cell):
    for paragraph in cell._element.xpath(".//w:drawing | .//w:pict"):
        return True
    return False

def transformer_tableau_et_images(doc_path):
    doc = Document(doc_path)
    nom_fichier = doc_path.name
    id_code = '-'.join(nom_fichier.replace('.docx', '').split('-')[:3])
    entete_inseree = False
    bloc_langue_actif = True
    langues_traitees = set()
    langue = None
    lignes_out = []

    for table in doc.tables:
        for row in table.rows:
            cellules = [cell.text.strip() for cell in row.cells]
            ligne = " | ".join(cellules).strip()
            if not ligne:
                continue

            image_in_cell = any(cell_contient_image(cell) for cell in row.cells)
            if image_in_cell:
                if ligne.startswith("#Q"):
                    ligne += " #IMGQ#"
                elif ligne.startswith("#A"):
                    ligne += " #IMGA#"

            if ligne.startswith("Défi |"):
                langue = 'fr'
                if langue in langues_traitees:
                    bloc_langue_actif = False
                    continue
                langues_traitees.add(langue)
                lignes_out.extend(generer_entete(langue, id_code, nom_fichier))
                entete_inseree = True
                bloc_langue_actif = True
                continue

            if ligne.startswith("Herausforderung |"):
                langue = 'de'
                if langue in langues_traitees:
                    bloc_langue_actif = False
                    continue
                langues_traitees.add(langue)
                lignes_out.extend(generer_entete(langue, id_code, nom_fichier))
                entete_inseree = True
                bloc_langue_actif = True
                continue

            if not bloc_langue_actif:
                continue

            if 'Temps passé à ce défi' in ligne or 'Für diese Herausforderung aufgewendete Zeit' in ligne:
                if entete_inseree:
                    lignes_out.append('##Work End')
                    lignes_out.append('')
                    if langue == 'fr':
                        lignes_out.append('[PAGEBREAK]')
                bloc_langue_actif = False
                continue

            lignes_out.extend(decouper_balises_speciales(ligne))

    lignes_out.append('')
    lignes_out.append('##Form End')

    doc_out = Document()
    previous_blank = False

    image_rels = [rel for rel in doc.part._rels.values() if rel.reltype == RT.IMAGE]

    # Vérification du nombre de balises #IMGQ# vs. images disponibles
    total_imgq = sum(l.count("#IMGQ#") for l in lignes_out)
    total_images = len(image_rels)

    print(f"[INFO] Balises #IMGQ# dans le document : {total_imgq}")
    print(f"[INFO] Images disponibles              : {total_images}")

    if total_imgq > total_images:
        print("[⚠️  AVERTISSEMENT] Il y a plus de balises #IMGQ# que d'images : certaines ne seront pas remplacées.")

    img_index = 0
    image_inserted = False

    for l in lignes_out:
        for sous_ligne in l.splitlines():
            if sous_ligne.strip() == '[PAGEBREAK]':
                doc_out.add_page_break()
                previous_blank = False
                continue

            # Gestion multiple balises #IMGQ#
            if "#IMGQ#" in sous_ligne:
                parts = sous_ligne.split("#IMGQ#")
                for i, part in enumerate(parts):
                    if part.strip():
                        p = doc_out.add_paragraph(part.strip())
                        p.paragraph_format.space_after = 0
                        p.paragraph_format.space_before = 0
                        p.paragraph_format.line_spacing = 1
                    if i < len(parts) - 1 and len(image_rels) > 0:
                        try:
                            img_data = image_rels[img_index].target_part.blob
                            image_name = image_rels[img_index].target_ref.rsplit("/", 1)[-1]
                            temp_path = OUTPUT_DIR / image_name
                            with open(temp_path, "wb") as img_file:
                                img_file.write(img_data)
                            doc_out.add_picture(str(temp_path), width=Inches(4.5))
                            print(f"[INFO] Image insérée à la place de #IMGQ# : {image_name}")
                            img_index = (img_index + 1) % len(image_rels)
                        except Exception as e:
                            doc_out.add_paragraph(f"[Erreur image : {str(e)}]")
                continue

            # Texte standard            
            if sous_ligne.strip():
                p = doc_out.add_paragraph(sous_ligne.strip())
                p.paragraph_format.space_after = 0
                p.paragraph_format.space_before = 0
                p.paragraph_format.line_spacing = 1
                previous_blank = False
            else:
                if not previous_blank:
                    p = doc_out.add_paragraph('')
                    p.paragraph_format.space_after = 0
                    p.paragraph_format.space_before = 0
                    p.paragraph_format.line_spacing = 1
                    previous_blank = True

    output_path = OUTPUT_DIR / nom_fichier
    doc_out.save(output_path)
    return nom_fichier

# Exécution principale
print("=== Informations sur les chemins ===")
print(f"Répertoire du script      : {base_dir}")
print(f"Dossier d'entrée          : {INPUT_DIR}")
print(f"Dossier de sortie         : {OUTPUT_DIR}")
print(f"Fichiers .docx détectés   : {len(files)}")
print("====================================\n")

for file in files:
    result = transformer_tableau_et_images(file)
    print(f"{file.name} : OK")
