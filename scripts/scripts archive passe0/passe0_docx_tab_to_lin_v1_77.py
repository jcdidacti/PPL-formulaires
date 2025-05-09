# ==============================================================================
# Script : passe0_docx_tab_to_lin_v1_76.py
# Objectif : insérer les images directement depuis les cellules, en respectant leur ordre réel
# Date : 2025-05-02
# ==============================================================================

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from datetime import datetime

SCRIPT_VERSION = "v1_77"

base_dir = Path(__file__).resolve().parent.parent
DATA_DIR = base_dir / "data"
INPUT_DIR = DATA_DIR / "00docx_tab"
OUTPUT_DIR = DATA_DIR / "01docx_lin_in"
LOG_DIR = OUTPUT_DIR / "log"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
LOG_DIR.mkdir(parents=True, exist_ok=True)

files = list(INPUT_DIR.glob("*.docx"))

def generer_entete(langue, id_code, nom_fichier):
    now = datetime.now().strftime('%Y-%m-%d %H:%M')
    return [
        '##Identification',
        f'#Script : {SCRIPT_VERSION}.py',
        f'#Run at : {now}',
        f'#ID file : {nom_fichier}',
        f'##LANG-{langue}',
        f'#ID : {id_code}',
        '#Version :',
        '#Date :',
        '#Author :',
        '',
        '##Introduction',
        '',
        '##Work Start'
    ]

def cell_contient_image(cell):
    for paragraph in cell._element.xpath(".//w:drawing | .//w:pict"):
        return True
    return False

def extraire_images_des_cellules(doc, output_dir):
    images = []
    index = 1
    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells[:2]):
                for drawing in cell._element.xpath(".//w:drawing"):
                    blips = drawing.xpath(".//a:blip", namespaces={"a": "http://schemas.openxmlformats.org/drawingml/2006/main"})
                    for blip in blips:
                        rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                        if rId:
                            rel = doc.part.related_parts.get(rId)
                            if rel:
                                image_data = rel.blob
                                image_name = f"image_{index:03d}_{'Q' if i == 0 else 'R'}.jpeg"
                                image_path = output_dir / image_name
                                with open(image_path, "wb") as f:
                                    f.write(image_data)
                                images.append((image_name, image_path))
                                index += 1
    return images

def transformer_tableau_et_images(doc_path):
    doc = Document(doc_path)
    nom_fichier = doc_path.name
    id_code = '-'.join(nom_fichier.replace('.docx', '').split('-')[:3])
    entete_inseree = False
    bloc_langue_actif = True
    langues_traitees = set()
    langue = None
    lignes_out = []

    doc_out = Document()
    previous_blank = False

    images_utilisees = extraire_images_des_cellules(doc, OUTPUT_DIR)

    for table in doc.tables:
        for row in table.rows:
            cellules = [cell.text.strip() for cell in row.cells]
            ligne = " | ".join(cellules).strip()
            if not ligne:
                continue

            if ligne.startswith("Défi |"):
                langue = 'fr'
                if langue in langues_traitees:
                    bloc_langue_actif = False
                    continue
                langues_traitees.add(langue)
                lignes_out.extend(generer_entete(langue, id_code, nom_fichier))
                entete_inseree = True
                bloc_langue_actif = True
                continue

            if ligne.startswith("Herausforderung |"):
                langue = 'de'
                if langue in langues_traitees:
                    bloc_langue_actif = False
                    continue
                langues_traitees.add(langue)
                lignes_out.extend(generer_entete(langue, id_code, nom_fichier))
                entete_inseree = True
                bloc_langue_actif = True
                continue

            if not bloc_langue_actif:
                continue

            if 'Temps passé à ce défi' in ligne or 'Für diese Herausforderung aufgewendete Zeit' in ligne:
                if entete_inseree:
                    lignes_out.append('##Work End')
                    lignes_out.append('')
                    if langue == 'fr':
                        lignes_out.append('[PAGEBREAK]')
                bloc_langue_actif = False
                continue

            lignes_out.append(ligne)

    lignes_out.append('')
    lignes_out.append('##Form End')

    for l in lignes_out:
        for sous_ligne in l.splitlines():
            if sous_ligne.strip() == '[PAGEBREAK]':
                doc_out.add_page_break()
                previous_blank = False
                continue
            if sous_ligne.strip():
                p = doc_out.add_paragraph(sous_ligne.strip())
                p.paragraph_format.space_after = 0
                p.paragraph_format.space_before = 0
                p.paragraph_format.line_spacing = 1
                previous_blank = False
            else:
                if not previous_blank:
                    p = doc_out.add_paragraph('')
                    p.paragraph_format.space_after = 0
                    p.paragraph_format.space_before = 0
                    p.paragraph_format.line_spacing = 1
                    previous_blank = True

    # Insertion des images détectées, pour contrôle visuel
    if images_utilisees:
        doc_out.add_page_break()
        doc_out.add_paragraph("Images extraites (ordre réel d'apparition dans le tableau) :")
        for image_name, image_path in images_utilisees:
            doc_out.add_picture(str(image_path), width=Inches(4.5))
            p = doc_out.add_paragraph(f"[image: {image_name}]")
            p.paragraph_format.space_after = 0
            p.paragraph_format.space_before = 0
            p.paragraph_format.line_spacing = 1

    output_path = OUTPUT_DIR / nom_fichier
    doc_out.save(output_path)
    return nom_fichier

# Exécution principale
print("=== Informations sur les chemins ===")
print(f"Répertoire du script      : {base_dir}")
print(f"Dossier d'entrée          : {INPUT_DIR}")
print(f"Dossier de sortie         : {OUTPUT_DIR}")
print(f"Fichiers .docx détectés   : {len(files)}")
print("====================================\n")

for file in files:
    result = transformer_tableau_et_images(file)
    print(f"{file.name} : OK")
