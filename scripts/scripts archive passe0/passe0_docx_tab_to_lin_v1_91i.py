# ==============================================================================
# Script : passe0_docx_tab_to_lin_v1_91i.py
# Objectif : nettoyage complet des \r/\n (↵), structuration correcte et élimination des lignes vides multiples
# Date : 2025-05-03
# ==============================================================================

import os
import hashlib
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from datetime import datetime

SCRIPT_VERSION = "v1_91i-debug"

base_dir = Path(__file__).resolve().parent.parent
DATA_DIR = base_dir / "data"
INPUT_DIR = DATA_DIR / "00docx_tab"
OUTPUT_DIR = DATA_DIR / "01docx_lin_in"
LOG_DIR = OUTPUT_DIR / "log"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
LOG_DIR.mkdir(parents=True, exist_ok=True)

files = list(INPUT_DIR.glob("*.docx"))

def generer_entete(langue, id_code, nom_fichier):
    now = datetime.now().strftime('%Y-%m-%d %H:%M')
    return [
        '##Identification',
        f'#Script : {SCRIPT_VERSION}.py',
        f'#Run at : {now}',
        f'#ID file : {nom_fichier}',
        f'##LANG-{langue}',
        f'#ID : {id_code}',
        '#Version :',
        '#Date :',
        '#Author :',
        '',
        '##Introduction',
        '',
        '##Work Start'
    ]

def extraire_images_des_cellules(doc, output_dir):
    images = []
    images_par_cellule = {}
    hash_dict = {}
    index = 1
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for i, cell in enumerate(row.cells[:2]):
                cle = (table_idx + 1, row_idx + 1, 'Q' if i == 0 else 'R')
                for drawing in cell._element.xpath(".//w:drawing"):
                    blips = drawing.xpath(".//*[local-name()='blip']")
                    for blip in blips:
                        rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                        if rId:
                            rel = doc.part.related_parts.get(rId)
                            if rel:
                                image_data = rel.blob
                                image_hash = hashlib.sha1(image_data).hexdigest()
                                if image_hash in hash_dict:
                                    image_name = hash_dict[image_hash]
                                    print(f"[DEBUG] Image déjà connue, réutilisée : {image_name} (cellule {cle})")
                                else:
                                    image_name = f"image_{index:03d}_{cle[2]}_T{cle[0]}R{cle[1]}.jpeg"
                                    image_path = output_dir / image_name
                                    with open(image_path, "wb") as f:
                                        f.write(image_data)
                                    hash_dict[image_hash] = image_name
                                    index += 1
                                    print(f"[DEBUG] Image extraite : {image_name} (table {cle[0]}, ligne {cle[1]}, cellule {cle[2]})")
                                images.append((image_name, output_dir / image_name))
                                images_par_cellule.setdefault(cle, []).append(image_name)
    print(f"[DEBUG] Total images enregistrées (sans doublons) : {len(hash_dict)}")
    print("[DEBUG] Dictionnaire images_par_cellule :")
    for cle, noms in images_par_cellule.items():
        print(f"  {cle} : {noms}")
    return images

def nettoyer_lignes(lignes):
    nettoyees = []
    ligne_vide_precedente = False
    for ligne in lignes:
        if ligne.strip() == "":
            if not ligne_vide_precedente:
                nettoyees.append("")
                ligne_vide_precedente = True
        else:
            nettoyees.append(ligne.strip())
            ligne_vide_precedente = False
    return nettoyees

def transformer_tableau_et_images(doc_path):
    doc = Document(doc_path)
    nom_fichier = doc_path.name
    id_code = '-'.join(nom_fichier.replace('.docx', '').split('-')[:3])
    entete_inseree = False
    bloc_langue_actif = True
    langues_traitees = set()
    langue = None
    lignes_out = []
    print("=== TEST POINT 1===")
    doc_out = Document()
    previous_blank = False

    extraire_images_des_cellules(doc, OUTPUT_DIR)

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            cellules = [cell.text.strip() for cell in row.cells]
            ligne = " | ".join(cellules).strip()
            if not ligne:
                continue

            if ligne.startswith("Défi |"):
                langue = 'fr'
                if langue in langues_traitees:
                    bloc_langue_actif = False
                    continue
                langues_traitees.add(langue)
                lignes_out.extend(generer_entete(langue, id_code, nom_fichier))
                entete_inseree = True
                bloc_langue_actif = True
                continue

            if ligne.startswith("Herausforderung |"):
                langue = 'de'
                if langue in langues_traitees:
                    bloc_langue_actif = False
                    continue
                langues_traitees.add(langue)
                lignes_out.extend(generer_entete(langue, id_code, nom_fichier))
                entete_inseree = True
                bloc_langue_actif = True
                continue

            if not bloc_langue_actif:
                continue

            if 'Temps passé à ce défi' in ligne or 'Für diese Herausforderung aufgewendete Zeit' in ligne:
                if entete_inseree:
                    lignes_out.append('##Work End')
                    lignes_out.append('')
                    if langue == 'fr':
                        lignes_out.append('[PAGEBREAK]')
                bloc_langue_actif = False
                continue

            sous_lignes = re.split(r'[\r\n]+', ligne)
            last_was_blank = False
            for sous_ligne in sous_lignes:
                sous_ligne = sous_ligne.strip()
                if sous_ligne:
                    lignes_out.append(sous_ligne)
                    last_was_blank = False
                else:
                    if not last_was_blank:
                        lignes_out.append("")
                        last_was_blank = True

    lignes_out.append('')
    lignes_out.append('##Form End')
    print("=== TEST POINT 3 ===")
    lignes_out = nettoyer_lignes(lignes_out)

    for ligne in lignes_out:
        if ligne.strip() == "[PAGEBREAK]":
            doc_out.add_page_break()
            previous_blank = False
        elif ligne.strip() == "":
            if not previous_blank:
                p = doc_out.add_paragraph('')
                p.paragraph_format.space_after = 0
                p.paragraph_format.space_before = 0
                p.paragraph_format.line_spacing = 1
                previous_blank = True
        else:
            p = doc_out.add_paragraph(ligne.strip())
            p.paragraph_format.space_after = 0
            p.paragraph_format.space_before = 0
            p.paragraph_format.line_spacing = 1
            previous_blank = False

    output_path = OUTPUT_DIR / nom_fichier
    doc_out.save(output_path)
    return nom_fichier

print("=== Informations sur les chemins ===")
print(f"Répertoire du script      : {base_dir}")
print(f"Dossier d'entrée          : {INPUT_DIR}")
print(f"Dossier de sortie         : {OUTPUT_DIR}")
print(f"Fichiers .docx détectés   : {len(files)}")
print("====================================\n")

for file in files:
    result = transformer_tableau_et_images(file)
    print(f"{file.name} : OK")
