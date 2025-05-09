# ==============================================================================
# Script : passe1_tableaux_to_linéaires_v1_02.py
# Objectif : Transformer des fichiers .docx avec tableaux en .docx linéaires
#            avec insertion de toutes les images à la fin
# Version : 1_02
# Date : 2025-04-11
# ==============================================================================

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# Répertoires dynamiques
base_dir = Path(os.path.abspath(__file__)).resolve().parent.parent
DATA_DIR = base_dir / "data"

INPUT_DIR = DATA_DIR / "00docx_tab"
OUTPUT_DIR = DATA_DIR / "01docx_lin_in"
LOG_DIR = OUTPUT_DIR / "log"

# Création des dossiers
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
LOG_DIR.mkdir(parents=True, exist_ok=True)

# Affichage des chemins
print("=== Informations sur les chemins ===")
print(f"Répertoire du script      : {base_dir.resolve()}")
print(f"Dossier d'entrée          : {INPUT_DIR.resolve()}")
print(f"Dossier de sortie         : {OUTPUT_DIR.resolve()}")



# Fonction pour séparer les balises #Q####### ou #A------- en début de ligne
def decouper_balises_speciales(texte):
    patterns = [r'^(#Q#######)(.*)', r'^(#A-------)(.*)']
    for pattern in patterns:
        match = re.match(pattern, texte.strip())
        if match:
            return [match.group(1), match.group(2).strip()]
    return [texte]
# Fonction pour séparer les balises #Q####### ou #A------- même après un '|'
def decouper_balises_speciales(texte):
    import re
    texte = texte.strip()
    if texte.startswith("|"):
        texte = texte[1:].strip()
    patterns = [
        r'^(#Q#######)(.*)',
        r'^(#A-------)(.*)',
        r'^.*\|\s*(#Q#######)(.*)',
        r'^.*\|\s*(#A-------)(.*)'
    ]
    for pattern in patterns:
        match = re.match(pattern, texte)
        if match:
            return [match.group(1), match.group(2).strip()]
    return [texte]
files = list(INPUT_DIR.glob("*.docx"))
print(f"Fichiers .docx détectés   : {len(files)}")
print("====================================\n")

def transformer_tableau_et_images(doc_path):
    doc = Document(doc_path)
    lignes = []
    nom_fichier = doc_path.name

    for table in doc.tables:
        for row in table.rows:
            cellules = [cell.text.strip() for cell in row.cells]
            lignes.append(" | ".join(cellules))

    doc_out = Document()
    for ligne in lignes:
        for item in decouper_balises_speciales(ligne):
            doc_out.add_paragraph(item)

    part = doc.part
    image_rels = [rel for rel in part._rels.values() if rel.reltype == RT.IMAGE]

    if image_rels:
        doc_out.add_paragraph("--- Images extraites du document source ---")

    image_count = 0
    for rel in image_rels:
        try:
            img_data = rel.target_part.blob
            image_name = rel.target_ref.rsplit("/", 1)[-1]
            temp_path = OUTPUT_DIR / image_name
            with open(temp_path, "wb") as img_file:
                img_file.write(img_data)
            doc_out.add_picture(str(temp_path), width=Inches(4.5))
            doc_out.add_paragraph(f"[image: {image_name}]")
            image_count += 1
        except Exception as e:
            doc_out.add_paragraph(f"[Erreur image : {str(e)}]")

    output_path = OUTPUT_DIR / nom_fichier
    doc_out.save(output_path)

    # Log de traitement
    log_path = LOG_DIR / nom_fichier.replace(".docx", ".log")
    with log_path.open("w", encoding="utf-8") as f:
        f.write(f"# Script : passe1_tableaux_to_linéaires_v1_02.py\n")
        f.write(f"# Version : 1_02\n")
        f.write(f"# Date : 2025-04-11\n")
        f.write("# ==========================================\n")
        f.write(f"Fichier transformé : {nom_fichier}\n")
        f.write(f"Lignes extraites : {len(lignes)}\n")
        f.write(f"Images extraites : {image_count}\n")

    return nom_fichier

# Exécution
log_global_path = LOG_DIR / "_global_transformation.log"
with log_global_path.open("w", encoding="utf-8") as f_global:
    f_global.write(f"# Script : passe1_tableaux_to_linéaires_v1_02.py\n")
    f_global.write(f"# Version : 1_02\n")
    f_global.write(f"# Date : 2025-04-11\n")
    f_global.write("# ==========================================\n")
    for file in files:
        result = transformer_tableau_et_images(file)
        f_global.write(f"{result} : OK\n")
        print(f"{result} : OK")
