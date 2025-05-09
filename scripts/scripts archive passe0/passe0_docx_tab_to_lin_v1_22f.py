# ==============================================================================
# Script : passe0_docx_tab_to_lin_v1_22c.py
# Objectif : Transformer des fichiers .docx avec tableaux en .docx linéaires
#            avec insertion de toutes les images à la fin
# Version : 1_02
# Date : 2025-04-11
# ==============================================================================

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from datetime import datetime

# Fonction pour générer l'en-tête structuré
def generer_entete(langue, id_code, nom_fichier):
    now = datetime.now().strftime('%Y-%m-%d %H:%M')
    return [
        '##Identification',
        '#Script : passe0_docx_tab_to_lin_v1_22c.py',
        f'#Run at : {now}',
        f'#ID file : {nom_fichier}',
        f'##LANG-{langue}',
        f'#ID : {id_code}',
        '#Version : ',
        '#Date : ',
        '#Author : ',
        '',
        '##Introduction',
        '',
        '##Work Start'
    ]

# Répertoires dynamiques
base_dir = Path(os.path.abspath(__file__)).resolve().parent.parent
DATA_DIR = base_dir / "data"

INPUT_DIR = DATA_DIR / "00docx_tab"
OUTPUT_DIR = DATA_DIR / "01docx_lin_in"
LOG_DIR = OUTPUT_DIR / "log"

# Création des dossiers
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
LOG_DIR.mkdir(parents=True, exist_ok=True)

# Affichage des chemins
print("=== Informations sur les chemins ===")
print(f"Répertoire du script      : {base_dir.resolve()}")
print(f"Dossier d'entrée          : {INPUT_DIR.resolve()}")
print(f"Dossier de sortie         : {OUTPUT_DIR.resolve()}")



# Fonction pour séparer les balises #Q####### ou #A------- en début de ligne
def decouper_balises_speciales(texte):
    patterns = [r'^(#Q#######)(.*)', r'^(#A-------)(.*)']
    for pattern in patterns:
        match = re.match(pattern, texte.strip())
        if match:
            return [match.group(1), match.group(2).strip()]
    return [texte]
# Fonction pour insérer un saut de ligne juste avant toute séquence #A-
def decouper_balises_speciales(texte):
    texte = texte.strip()
    if '#A-' in texte:
        idx = texte.find('#A-')
        return [texte[:idx].strip(), texte[idx:].strip()]
    return [texte]
files = list(INPUT_DIR.glob("*.docx"))
print(f"Fichiers .docx détectés   : {len(files)}")
print("====================================\n")

def transformer_tableau_et_images(doc_path):
    doc = Document(doc_path)
    lignes = []
    nom_fichier = doc_path.name

    for table in doc.tables:
        for row in table.rows:
            cellules = [cell.text.strip() for cell in row.cells]
            lignes.append(" | ".join(cellules))

    doc_out = Document()

# Exécution
log_global_path = LOG_DIR / "_global_transformation.log"
with log_global_path.open("w", encoding="utf-8") as f_global:
    f_global.write(f"# Script : passe0_docx_tab_to_lin_v1_22d.py\n")
    f_global.write(f"# Version : 1_22d\n")
    f_global.write(f"# Date : {datetime.now().strftime('%Y-%m-%d')}\n")
    f_global.write("# ==========================================\n")
    for file in files:
        result = transformer_tableau_et_images(file)
        f_global.write(f"{result} : OK\n")
        print(f"{result} : OK")
