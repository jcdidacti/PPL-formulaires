
# ==============================================================================
# Script : passe1_docx_lin_to_txt_v2_69e.py
# Objectif : Gérer plusieurs images dans un même paragraphe (rel_id multiples)
# Date : 2025-04-20
# ==============================================================================

import os
from pathlib import Path
from docx import Document
from hashlib import md5
from datetime import datetime
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# === INIT RÉPERTOIRES ===
base_dir = Path(__file__).resolve().parent.parent
data_dir = base_dir / "data"

input_dir = data_dir / "02docx_lin_out"
output_dir = data_dir / "02text_p1_out"
image_dir = output_dir / "images"
log_dir = output_dir / "log"

output_dir.mkdir(parents=True, exist_ok=True)
image_dir.mkdir(parents=True, exist_ok=True)
log_dir.mkdir(parents=True, exist_ok=True)

global_log_path = log_dir / "_global_extraction.log"
global_log_entries = []

# === STRUCTURE HEADER/FOOTER ===
def get_struct(nom_fichier, code, auteur):
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    header = [
        "## Identification",
        f"#Script : passe1_docx_lin_to_txt_v2_69e.py",
        f"#Run at : {now}",
        f"#ID file : {nom_fichier}",
        f"ID        {code}",
        f"Version   00.00",
        f"Date      20-01-2025",
        f"Auteur    {auteur}",
        "",
        "## Introduction",
        "# To be completed",
        "",
        "## Work Start"
    ]
    footer = [
        "",
        "## Work Stop",
        "",
        "## Personal data",
        "Temps passé sur ce questionnaire :",
        "Code de l'élève :",
        "Code de vérification :",
        "## End of Form"
    ]
    return header, footer

def process_docx(docx_path: Path):
    nom_fichier = docx_path.stem
    code = "-".join(nom_fichier.split("-")[:3])
    auteur = "MC"
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    print(f"🔍 Traitement : {nom_fichier}")

    doc = Document(docx_path)

    image_hash_map = {}
    image_hash_index = {}
    relid_to_hash = {}
    image_counter = 1
    bloc_texte = []

    try:
        print("📥 Extraction des images...")
        for rel_id, rel in doc.part._rels.items():
            if rel.reltype == RT.IMAGE:
                img_data = rel.target_part.blob
                img_hash = md5(img_data).hexdigest()
                relid_to_hash[rel_id] = img_hash
                if img_hash not in image_hash_map:
                    ext = rel.target_part.content_type.split("/")[-1]
                    img_name = f"{code}_image{image_counter:03}.{ext}"
                    img_path = image_dir / img_name
                    with open(img_path, "wb") as f:
                        f.write(img_data)
                    image_hash_map[img_hash] = img_name
                    image_hash_index[img_hash] = image_counter
                    print(f"✅ Image enregistrée : {img_name}")
                    image_counter += 1
                else:
                    print(f"🔁 Image déjà connue : {image_hash_map[img_hash]}")

        print("📄 Lecture des paragraphes...")
        for para in doc.paragraphs:
            drawings = para._element.xpath('.//w:drawing')
            text = para.text.strip()
            has_image = bool(drawings)

            if "défi" in text.lower() and "|" in text:
                parts = text.split("|")
                if len(parts) == 2 and parts[1].strip():
                    auteur = parts[1].strip()
                    continue

            if "|" in text:
                q, a = text.split("|", 1)
                bloc_texte.append("#Q######################################################")
                bloc_texte.append(q.strip())
                bloc_texte.append("#A------------------------------------------------------")
                bloc_texte.append(f"#-[type:text] {a.strip()} -#")
            elif text:
                bloc_texte.append(text)

            if has_image:
                rels = para._element.xpath('.//a:blip/@r:embed')
                for rel_id in rels:
                    img_hash = relid_to_hash.get(rel_id, None)
                    if img_hash and img_hash in image_hash_map:
                        img_name = image_hash_map[img_hash]
                        img_index = image_hash_index[img_hash]
                        bloc_texte.append(f"#PICT{img_index:03}# [image: {img_name}]")
                        print(f"🖋️ Insertion : #PICT{img_index:03}# [image: {img_name}]")
                    else:
                        bloc_texte.append("#PICT000# [image: image_inconnue.png]")
                        print(f"⚠️ Image inconnue pour rel_id {rel_id}")

        header, footer = get_struct(nom_fichier, code, auteur)
        final_content = header + bloc_texte + footer + ["", "## Images"]

        txt_path = output_dir / f"{nom_fichier}.txt"
        txt_path.write_text("\n".join(final_content), encoding="utf-8")
        print(f"✅ Fichier écrit : {txt_path}")

        log_path = log_dir / f"{nom_fichier}.log"
        with open(log_path, "w", encoding="utf-8") as logf:
            logf.write(f"# Script : passe1_docx_lin_to_txt_v2_69e.py\n")
            logf.write(f"# Execution Time : {timestamp}\n")
            logf.write(f"# File processed : {docx_path.name}\n")
            logf.write(f"# Paragraphs : {len(doc.paragraphs)}\n")
            logf.write(f"# Images extraites : {len(image_hash_map)}\n")

        global_log_entries.append(f"{nom_fichier} : OK")
        return True

    except Exception as e:
        print(f"❌ Erreur : {e}")
        global_log_entries.append(f"{nom_fichier} : NOT OK - {e}")
        return False

if __name__ == "__main__":
    fichiers = list(input_dir.glob("*.docx"))
    for fichier in fichiers:
        if fichier.name.startswith("~$"):
            print(f"⏭️ Fichier ignoré : {fichier.name}")
            continue
        process_docx(fichier)

    with open(global_log_path, "w", encoding="utf-8") as f:
        f.write("# Global log — passe1 v2.69e\n")
        f.write(f"# Run : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write("# ===========================================\n")
        for entry in global_log_entries:
            f.write(entry + "\n")
