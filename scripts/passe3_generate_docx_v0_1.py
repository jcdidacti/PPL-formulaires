
# ==============================================================================
# Script : passe3_generate_docx_v0_1.py
# Objectif : Générer les fichiers instructeur et élève à partir d’un .txt structuré
# Date : Phase prototype
# ==============================================================================

from docx import Document
from docx.shared import Pt, Cm
import re
from pathlib import Path
from datetime import datetime

# === CONFIGURATION ===
TITRE_STYLE = "Titre 1"
QUESTION_STYLE = "Heading 2"
REPONSE_STYLE_INSTR = "Normal"
REPONSE_STYLE_ELEVE = "ChampÉlève"  # à créer dans le modèle si souhaité
IMAGE_DEFAULT_HEIGHT_CM = 5

# === Fonction d'extraction des blocs de question ===
def extraire_blocs_questions(txt_path):
    with open(txt_path, encoding="utf-8") as f:
        lignes = f.readlines()

    blocs = []
    meta = {"id": "", "auteur": "", "version": "", "date": ""}
    langue = None
    i = 0
    while i < len(lignes):
        ligne = lignes[i].strip()

        # Métadonnées
        if ligne.startswith("ID"):
            meta["id"] = ligne.split()[1]
        elif ligne.startswith("Auteur"):
            meta["auteur"] = ligne.split(":", 1)[1].strip()
        elif ligne.startswith("Version"):
            meta["version"] = ligne.split()[1]
        elif ligne.startswith("Date"):
            meta["date"] = ligne.split()[1]

        elif ligne.startswith("##LANG-"):
            langue = ligne.replace("##LANG-", "").strip()

        elif ligne.startswith("#Q"):
            question = lignes[i + 1].strip()
            ref = ""
            reponse = ""
            image = None
            if lignes[i + 2].strip().startswith("#A"):
                i += 2
                contenu = lignes[i + 1].strip()
                ref_match = re.match(r"(BAK\s?[\d\.\-\s]+):?", contenu)
                if ref_match:
                    ref = ref_match.group(1).strip()
                    contenu = contenu.replace(ref, "").lstrip(" :")
                reponse = contenu
            blocs.append({
                "question": question,
                "reponse": reponse,
                "ref": ref,
                "langue": langue
            })
        i += 1
    return meta, blocs

# === Prototype génération document ===
def generer_document(blocs, meta, output_path, instructeur=True):
    doc = Document()
    doc.add_heading("Document " + ("Instructeur" if instructeur else "Élève"), 0)

    doc.add_paragraph(f"ID        : {meta['id']}")
    doc.add_paragraph(f"Auteur    : {meta['auteur']}")
    doc.add_paragraph(f"Version   : {meta['version']}")
    doc.add_paragraph(f"Date      : {meta['date']}")
    doc.add_paragraph(f"Généré le : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")

    for i, bloc in enumerate(blocs, 1):
        doc.add_paragraph(f"[Q{i:02}] {bloc['question']}", style=QUESTION_STYLE)
        if bloc["ref"]:
            doc.add_paragraph(f"📘 Référence : {bloc['ref']}", style="Intense Quote")

        if instructeur:
            doc.add_paragraph(f"Réponse attendue :", style="Normal")
            doc.add_paragraph(bloc["reponse"], style=REPONSE_STYLE_INSTR)
        else:
            doc.add_paragraph(f"Réponse à compléter :", style="Normal")
            doc.add_paragraph("__________________________________", style=REPONSE_STYLE_ELEVE)

        doc.add_paragraph("-" * 80)

    doc.save(output_path)

# === Exemple d’appel (à adapter en script principal)
if __name__ == "__main__":
    txt_input = Path("exemple_questionnaire.txt")
    meta, blocs = extraire_blocs_questions(txt_input)
    generer_document(blocs, meta, "exemple_instr.docx", instructeur=True)
    generer_document(blocs, meta, "exemple_eleve.docx", instructeur=False)
