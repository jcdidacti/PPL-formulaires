# ==============================================================================
# Script : passe1_docx_lin_to_txt_v2_73.py
# Objectif : Traitement complet multilingue avec structure, Q/R, images et langue
# Date : 2025-04-20
# ==============================================================================

import os
import re
from pathlib import Path
from docx import Document

def verifier_entete(doc):
    """Vérifie que les 5 premières balises d'en-tête sont présentes dans les 10 premières lignes."""
    attendues = ["##Identification", "#Script", "#Run at", "#ID file", "##LANG-"]
    lignes = [p.text.strip() for p in doc.paragraphs[:10] if p.text.strip()]
    print(f"[DEBUG] lignes analysées : {lignes}")
    correspondances = sum(any(ligne.startswith(att) for ligne in lignes) for att in attendues)
    for att in attendues:
        if not any(ligne.startswith(att) for ligne in lignes):
            print(f"[WARN] Balise manquante ou incorrecte : {att}")
    return correspondances == len(attendues)

from hashlib import md5
from datetime import datetime
from docx.opc.constants import RELATIONSHIP_TYPE as RT

import subprocess

import subprocess

def get_git_version():
    try:
        tag = subprocess.check_output(
            ["git", "describe", "--tags", "--exact-match"],
            stderr=subprocess.DEVNULL
        ).decode().strip()
        return tag
    except Exception:
        return "non tagué"

VERSION_SCRIPT = get_git_version()


# === Configuration ===
NB_LIGNES_DETECTION = 40

# === Mots typiques pour détection de langue ===
mots_typiques = {
    "fr": {"le", "la", "est", "une", "des", "avec", "ne", "pas", "dans", "ce",
           "quelle", "l’instrument", "altimètre", "indiquée", "pression", "réglée"},
    "de": {"und", "die", "ist", "nicht", "das", "mit", "auf", "ein", "der"},
    "it": {"è", "con", "una", "per", "non", "il", "nella", "che", "di"}
}

# === Répertoires ===
base_dir = Path(__file__).resolve().parent.parent
data_dir = base_dir / "data"
input_dir = data_dir / "02docx_lin_out"
output_dir = data_dir / "02text_p1_out"
image_dir = output_dir / "images"
log_dir = output_dir / "log"

output_dir.mkdir(parents=True, exist_ok=True)
image_dir.mkdir(parents=True, exist_ok=True)
log_dir.mkdir(parents=True, exist_ok=True)

global_log_path = log_dir / "_global_extraction.log"
global_log_entries = []

# === Détection de langue ===
def detect_lang(text):
    words = re.findall(r"\b\w+\b", text.lower())
    scores = {lang: sum(1 for w in words if w in mots) for lang, mots in mots_typiques.items()}
    best = max(scores, key=scores.get)
    return best if scores[best] > 2 else None

# === Création de l'en-tête pour une langue donnée
def make_header(code_langue, nom_fichier, auteur="MC"):
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    code = "-".join(nom_fichier.split("-")[:3])
    return [
        "## Identification",
        f"#Script : passe1_docx_lin_to_txt_v2_73.py",
        f"#Run at : {now}",
        f"#ID file : {nom_fichier}",
        f"##LANG-{code_langue}",
        f"ID        {code}",
        f"Version   00.00",
        f"Date      20-01-2025",
        f"Auteur    {auteur}",
        ""
    ]

# === Traitement principal
def process_docx(docx_path: Path):
    nom_fichier = docx_path.stem
    print(f"🔍 Traitement : {nom_fichier}", end="")
    doc = Document(docx_path)

    if not verifier_entete(doc):
        print("❌ Erreur : aucune balise détectée au début du document. Veuillez vérifier la structure initiale.")
        return

    langue_courante = None
    bloc_detecte = False
    bloc_texte = []
    image_hash_map = {}
    image_hash_index = {}
    relid_to_hash = {}
    image_counter = 1
    compteur_questions = {}
    compteur_images = {}

    try:
        # Extraction des images
        for rel_id, rel in doc.part._rels.items():
            if rel.reltype == RT.IMAGE:
                img_data = rel.target_part.blob
                img_hash = md5(img_data).hexdigest()
                relid_to_hash[rel_id] = img_hash
                if img_hash not in image_hash_map:
                    ext = rel.target_part.content_type.split("/")[-1]
                    img_name = f"{nom_fichier}_image{image_counter:03}.{ext}"
                    img_path = image_dir / img_name
                    with open(img_path, "wb") as f:
                        f.write(img_data)
                    image_hash_map[img_hash] = img_name
                    image_hash_index[img_hash] = image_counter
                    image_counter += 1

        for para in doc.paragraphs:
            texte = para.text.strip()
            if not texte:
                continue

            if texte.startswith("##LANG-"):
                match = re.match(r"##LANG-([a-zA-Z]{2})\b", texte)
                if not match:
                    raise ValueError(f"❌ Balise incorrecte : {texte}")
                code = match.group(1).lower()
                if code not in mots_typiques:
                    raise ValueError(f"❌ Balise inconnue : {texte}")
                langue_courante = code
                bloc_detecte = True
                bloc_texte += make_header(code, nom_fichier)
                compteur_questions[code] = 0
                compteur_images[code] = 0
                continue

            if not bloc_detecte:
                lignes = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                texte_bloc_initial = " ".join(lignes[:NB_LIGNES_DETECTION])
                langue_prob = detect_lang(texte_bloc_initial)
                suggestion = f"##LANG-{langue_prob}" if langue_prob else "langue indéterminée"
                print(f"❌ Erreur : aucune balise détectée au début du document. Probablement {langue_prob or '???'}. Veuillez insérer la balise {suggestion} au début.\n")
                global_log_entries.append(f"{nom_fichier} : NOT OK")
                return False

            langue_detectee = detect_lang(texte)
            if langue_detectee and langue_detectee != langue_courante:
                print(f"❌ Changement de langue détecté dans le texte. Probablement au passage {langue_courante} - {langue_detectee}. Veuillez insérer la balise ##LANG-{langue_detectee} au changement.\n")
                global_log_entries.append(f"{nom_fichier} : NOT OK")
                return False

            # Questions / réponses
            if "|" in texte:
                q, a = texte.split("|", 1)
                bloc_texte.append("#Q######################################################")
                bloc_texte.append(q.strip())
                bloc_texte.append("#A------------------------------------------------------")
                bloc_texte.append(f"#-[type:text] {a.strip()} -#")
                compteur_questions[langue_courante] += 1
            else:
                bloc_texte.append(texte)

            # Images
            rels = para._element.xpath('.//a:blip/@r:embed')
            for rel_id in rels:
                img_hash = relid_to_hash.get(rel_id)
                if img_hash and img_hash in image_hash_map:
                    img_name = image_hash_map[img_hash]
                    img_index = image_hash_index[img_hash]
                    bloc_texte.append(f"#PICT{img_index:03}# [image: {img_name}]")
                compteur_images[langue_courante] += 1

        # Ajouter le pied de page et écrire le fichier
        bloc_texte += ["", "## Work Stop", "", "## End of Form"]
        txt_path = output_dir / f"{nom_fichier}.txt"
        txt_path.write_text("\n".join(bloc_texte), encoding="utf-8")

        print("   ✅ Langue cohérente dans tout le document\n")

        log_path = log_dir / f"{nom_fichier}.log"
        with open(log_path, "w", encoding="utf-8") as logf:
            logf.write(f"# Log de traitement —  passe1 (version : {VERSION_SCRIPT})\n")
            logf.write(f"# Fichier     : {nom_fichier}\n")
            logf.write(f"# Date        : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            logf.write(f"# Résultat    : OK\n\n")
            logf.write("## Résumé par langue\n\n")
            logf.write("Langue   | Questions | Images\n")
            logf.write("-------- | ----------|--------\n")
            for lang in compteur_questions:
                q = compteur_questions[lang]
                i = compteur_images.get(lang, 0)
                logf.write(f"{lang:<8} | {q:<9} | {i}\n")

        global_log_entries.append(f"{nom_fichier} : OK")
        return True

    except Exception as e:
        print(f"❌ Erreur critique : {e}\n")

        log_path = log_dir / f"{nom_fichier}.log"
        with open(log_path, "w", encoding="utf-8") as logf:
            logf.write(f"# Log de traitement — passe1 (version : {VERSION_SCRIPT})\n")
            logf.write(f"# Fichier     : {nom_fichier}\n")
            logf.write(f"# Date        : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            logf.write(f"# Résultat    : NOT OK\n\n")
            logf.write(f"{e}\n")

        global_log_entries.append(f"{nom_fichier} : NOT OK")
        return False

if __name__ == "__main__":
    fichiers = list(input_dir.glob("*.docx"))
    for fichier in fichiers:
        if fichier.name.startswith("~$"):
            continue
        process_docx(fichier)

    with open(global_log_path, "w", encoding="utf-8") as f:
        f.write(f"# Global log — passe1 (version : {VERSION_SCRIPT})\n")
        f.write(f"# Run : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write("# ===========================================\n")
        for entry in global_log_entries:
            f.write(entry + "\n")
